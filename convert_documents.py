import os
import re
from docx import Document
from doc2docx import convert

# Function to handle the conversion logic dynamically
def convert_line(line):
    # Rule 1: Replace BV [MT05] with BV {{caseReference}}
    line = re.sub(r'BV \[MT05\]', 'BV {{caseReference}}', line)

    # Rule 2: Replace specific combined placeholders with caseReference
    if re.search(r'\[MT05\](?:.*\[MT11|\b)', line):  # Matches [MT05] or [MT05] with additional [MT11(2)]
        line = re.sub(r'\[MT05\](?:.*\[MT11\(.*?\)\])?', '{{caseReference}}', line)
    
    # Rule 3: Remove diary-related lines
    if re.search(r'\[&Diary|\[&HISTORY', line):
        return ""  # Empty string to remove the line
    
    # Rule 4: Handle selective placeholders and ignore specific tags
    line = re.sub(r'\[&Message.*?\]', '', line)  # Remove [&Message ...]
    line = re.sub(r'\[&Show.*?\]', '', line)  # Remove [&Show ...]
    
    # Rule 5: Replace DATE placeholders with {{currentDate}}
    line = re.sub(r'\[DATE:DS\(".*?"\)\]', '{{currentDate}}', line)
    
    # Replace remaining placeholders with dynamic Apryse equivalents
    def convert_placeholder(match):
        placeholder = match.group(1)
        # General rule: remove special characters and wrap in curly braces
        cleaned_placeholder = re.sub(r'[^\w]', '', placeholder)
        return "{{" + cleaned_placeholder + "}}"
    
    return re.sub(r'\[(.*?)\]', convert_placeholder, line)

# Function to process each DOCX file
def process_docx(file_path):
    doc = Document(file_path)

    # Process each paragraph and apply the conversion
    for paragraph in doc.paragraphs:
        paragraph.text = convert_line(paragraph.text)

    # Handle tables if necessary
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = convert_line(cell.text)

    # Save the modified document with a new name
    output_path = os.path.join(os.path.dirname(file_path), "output_" + os.path.basename(file_path))
    doc.save(output_path)

# Directory containing the .doc and .docx files (current directory)
directory = os.getcwd()

# Step 1: Convert all .doc and .DOC files to .docx
for filename in os.listdir(directory):
    if filename.lower().endswith(".doc") and not filename.startswith("output"):
        file_path = os.path.join(directory, filename)
        convert(file_path)  # Convert .doc or .DOC to .docx

# Step 2: Process all .docx files
for filename in os.listdir(directory):
    if filename.endswith(".docx") and not filename.startswith("output"):
        file_path = os.path.join(directory, filename)
        process_docx(file_path)